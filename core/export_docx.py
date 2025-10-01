from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

def export_docx_bytes(title, exec_summary, kpi_dict, chart_png_bytes=None, logo_bytes=None):
    doc = Document()

    # Optional logo
    if logo_bytes:
        try:
            doc.add_picture(io.BytesIO(logo_bytes), width=Inches(2.0))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.LEFT
        except Exception:
            pass

    h = doc.add_heading(title, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT

    p = doc.add_paragraph(exec_summary or "")
    p.paragraph_format.space_after = Pt(12)

    doc.add_heading("KPIs", level=2)
    if kpi_dict:
        for k, v in kpi_dict.items():
            doc.add_paragraph(f"• {k}: {v}")

    if chart_png_bytes:
        doc.add_heading("Chart", level=2)
        try:
            doc.add_picture(io.BytesIO(chart_png_bytes), width=Inches(5.5))
        except Exception:
            pass

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
